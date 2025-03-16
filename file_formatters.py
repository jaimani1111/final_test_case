# file_formatters.py
from docx import Document
from io import BytesIO
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils.dataframe import dataframe_to_rows

def format_test_cases_for_txt(test_cases):
    """Format test cases for TXT download"""
    formatted_text = ""
    for case in test_cases:
        formatted_text += f"Test Case {case['Sl No.']}: {case['Test Case Description']}\n"
        formatted_text += f"Sl No.: {case['Sl No.']}\n"
        formatted_text += f"Requirement ID: {case['Requirement ID']}\n"
        formatted_text += f"Test Case ID: {case['Test Case ID']}\n"
        formatted_text += f"Module: {case['Module']}\n"
        formatted_text += f"Test Case Description: {case['Test Case Description']}\n"
        formatted_text += f"Execution Steps:\n{case['Execution Steps']}\n"
        formatted_text += f"Expected Result: {case['Expected Result']}\n\n"
    return formatted_text

def format_test_cases_for_docx(test_cases):
    """Format test cases for Word (DOCX) download"""
    doc = Document()
    for case in test_cases:
        doc.add_heading(f"Test Case {case['Sl No.']}: {case['Test Case Description']}", level=1)
        doc.add_paragraph(f"Sl No.: {case['Sl No.']}")
        doc.add_paragraph(f"Requirement ID: {case['Requirement ID']}")
        doc.add_paragraph(f"Test Case ID: {case['Test Case ID']}")
        doc.add_paragraph(f"Module: {case['Module']}")
        doc.add_paragraph(f"Test Case Description: {case['Test Case Description']}")
        doc.add_paragraph("Execution Steps:")
        doc.add_paragraph(case['Execution Steps'])
        doc.add_paragraph(f"Expected Result: {case['Expected Result']}")
        doc.add_paragraph("\n")  # Add space between test cases
    return doc

def format_test_cases_for_excel(test_cases):
    """Format test cases for Excel download"""
    excel_buffer = BytesIO()
    wb = Workbook()
    ws = wb.active
    
    # Add headers
    headers = ["Sl No.", "Requirement ID", "Test Case ID", "Module", "Test Case Description", "Execution Steps", "Expected Result"]
    ws.append(headers)
    
    # Add data
    for case in test_cases:
        ws.append([
            case['Sl No.'],
            case['Requirement ID'],
            case['Test Case ID'],
            case['Module'],
            case['Test Case Description'],
            case['Execution Steps'],
            case['Expected Result']
        ])
    
    # Format columns
    for col in ws.columns:
        max_length = max(len(str(cell.value)) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = max_length + 2
    
    # Highlight the header row in yellow
    yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
    for cell in ws[1]:  # Header row (first row)
        cell.fill = yellow_fill
    
    # Format steps
    for row in ws.iter_rows(min_row=2):
        cell = row[5]  # Execution Steps column
        cell.alignment = Alignment(wrap_text=True, vertical='top')
    
    wb.save(excel_buffer)
    return excel_buffer.getvalue()