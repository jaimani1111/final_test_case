import os
import pandas as pd
from dotenv import load_dotenv
from groq import Groq
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
import streamlit as st
from docx import Document
from io import BytesIO
import re
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils.dataframe import dataframe_to_rows

# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

groq_client = Groq(api_key=GROQ_API_KEY)
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Load Chroma DB
try:
    vector_db = Chroma(persist_directory="./chroma_db", embedding_function=embeddings)
except Exception as e:
    st.error(f"Error loading Chroma DB: {e}")
    st.stop()

def clean_ai_response(response):
    """Improved parser to handle variable numbers of execution steps"""
    test_cases = []
    
    # Split into test case blocks using numbered headers
    case_blocks = re.split(r'\n\*\*Test Case \d+:', response)
    
    for block in case_blocks[1:]:  # Skip first empty split
        current_case = {
            "Sl No.": "",
            "Requirement ID": "",
            "Test Case ID": "",
            "Module": "",
            "LOB": "",
            "Region": "",
            "Test Case Description": "",
            "Execution Steps": "",
            "Expected Result": ""
        }
        
        # Extract fields using regex
        sl_no_match = re.search(r'Sl No\.:\s*(\d+)', block)
        if sl_no_match:
            current_case["Sl No."] = sl_no_match.group(1).strip()
        
        req_id_match = re.search(r'Requirement ID:\s*([^\n]+)', block)
        if req_id_match:
            current_case["Requirement ID"] = req_id_match.group(1).strip()
        
        tc_id_match = re.search(r'Test Case ID:\s*([^\n]+)', block)
        if tc_id_match:
            current_case["Test Case ID"] = tc_id_match.group(1).strip()
        
        module_match = re.search(r'Module:\s*([^\n]+)', block)
        if module_match:
            current_case["Module"] = module_match.group(1).strip()
        
        lob_match = re.search(r'LOB:\s*([^\n]+)', block)
        if lob_match:
            current_case["LOB"] = lob_match.group(1).strip()
        
        region_match = re.search(r'Region:\s*([^\n]+)', block)
        if region_match:
            current_case["Region"] = region_match.group(1).strip()
        
        desc_match = re.search(r'Test Case Description:\s*([^\n]+)', block)
        if desc_match:
            current_case["Test Case Description"] = desc_match.group(1).strip()
        
        # Extract execution steps dynamically
        steps_match = re.search(r'Execution Steps:\s*(.*?)(?=\nExpected Result:)', block, re.DOTALL)
        if steps_match:
            steps = steps_match.group(1).strip()
            # Ensure each step is on a new line with a gap line between steps
            steps = re.sub(r'Step(\d+):', r'\nStep\1:', steps)  # Add newline before each step
            steps = steps.replace('\n', '\n\n')  # Add an extra newline to create a gap
            steps = steps.strip()  # Remove leading/trailing whitespace
            current_case["Execution Steps"] = steps
        
        # Extract expected result
        result_match = re.search(r'Expected Result:\s*(.*)', block, re.DOTALL)
        if result_match:
            current_case["Expected Result"] = result_match.group(1).strip()
        
        # Add only if all required fields are present
        if all(current_case.values()):
            test_cases.append(current_case)
        else:
            # Debugging: Print the block that failed to parse
            st.warning(f"Failed to parse block: {block}")
    
    return test_cases

def generate_test_cases(insurance_type, region, line_of_business, user_requirements, num_test_cases):
    """Generate test cases with variable execution steps"""
    query = f"{insurance_type} {region} {line_of_business} {user_requirements}"
    
    # Print the query used for similarity search
    print(f"Query for similarity search: {query}")
    
    # Perform similarity search
    similar_cases = vector_db.similarity_search(query, k=4)
    
    # Print the similar cases retrieved
    print("Similar cases retrieved from vector database:")
    for i, case in enumerate(similar_cases, 1):
        print(f"Case {i}:\n{case.page_content}\n{'-' * 50}")
    
    # Construct context from similar cases
    context = "\n".join([case.page_content for case in similar_cases])
    
    # Print the context used for generating new test cases
    print("Context constructed from similar cases:")
    print(context)
    
    # Construct the prompt
    prompt = f"""As an Insurance QA Expert, create {num_test_cases} test cases. Generated test cases should be based on the specific line of business and region. Please make sure the generated test cases are very high in quality and detail with this structure:

**Test Case X: [Scenario]**

Sl No.: [number]
Requirement ID: [REQ-XXX]
Test Case ID: [TC-XXX]
Module: [Specific Module]
LOB: {line_of_business}
Region: {region}
Test Case Description: [Clear description]
Execution Steps:
Step1: [Action]
Step2: [Action]
... (Add as many steps as needed for the test case)
Expected Result: [Measurable outcome]

Avoid these examples but take inspiration from them only: {context}"""
    
    # Print the prompt sent to the Groq API
    print("Prompt sent to Groq API:")
    print(prompt)
    
    # Send the prompt to Groq API
    response = groq_client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama3-70b-8192",
    )
    
    raw_content = response.choices[0].message.content
    st.session_state.raw_response = raw_content  # Store for debugging
    return clean_ai_response(raw_content)

def format_test_cases_for_txt(test_cases):
    """Format test cases for TXT download"""
    formatted_text = ""
    for case in test_cases:
        formatted_text += f"Test Case {case['Sl No.']}: {case['Test Case Description']}\n"
        formatted_text += f"Sl No.: {case['Sl No.']}\n"
        formatted_text += f"Requirement ID: {case['Requirement ID']}\n"
        formatted_text += f"Test Case ID: {case['Test Case ID']}\n"
        formatted_text += f"Module: {case['Module']}\n"
        formatted_text += f"Test Case Description: {case['Test Case Description']}\n"
        formatted_text += f"Execution Steps:\n{case['Execution Steps']}\n"
        formatted_text += f"Expected Result: {case['Expected Result']}\n\n"
    return formatted_text

def format_test_cases_for_docx(test_cases):
    """Format test cases for Word (DOCX) download"""
    doc = Document()
    for case in test_cases:
        doc.add_heading(f"Test Case {case['Sl No.']}: {case['Test Case Description']}", level=1)
        doc.add_paragraph(f"Sl No.: {case['Sl No.']}")
        doc.add_paragraph(f"Requirement ID: {case['Requirement ID']}")
        doc.add_paragraph(f"Test Case ID: {case['Test Case ID']}")
        doc.add_paragraph(f"Module: {case['Module']}")
        doc.add_paragraph(f"Test Case Description: {case['Test Case Description']}")
        doc.add_paragraph("Execution Steps:")
        doc.add_paragraph(case['Execution Steps'])
        doc.add_paragraph(f"Expected Result: {case['Expected Result']}")
        doc.add_paragraph("\n")  # Add space between test cases
    return doc


# Streamlit UI
st.markdown("<h1 style='text-align: center;'>XC AI Test Case Generator</h1>", unsafe_allow_html=True)

with st.sidebar:
    st.header("Configuration")
    insurance_type = st.selectbox("Insurance Type", ["Auto", "Health", "Home", "Life"])
    region = st.selectbox("Region", ["North America", "Europe", "Asia Pacific", "ANZ"])
    line_of_business = st.selectbox("Line of Business", ["Retail", "Commercial", "Enterprise"])
    
    st.subheader("Acceptance Criteria")
    
    # Initialize session state for file uploader and manual input
    if 'file_uploaded' not in st.session_state:
        st.session_state.file_uploaded = False
    if 'manual_input_used' not in st.session_state:
        st.session_state.manual_input_used = False
    
    # File uploader
    criteria_file = st.file_uploader("Upload File (txt/csv/docx)", type=["txt", "csv", "docx"], 
                                    disabled=st.session_state.manual_input_used)
    
    # If a file is uploaded, disable manual input
    if criteria_file:
        st.session_state.file_uploaded = True
        st.session_state.manual_input_used = True
    else:
        st.session_state.file_uploaded = False
    
    # Manual input text area
    manual_input = st.text_area("Or enter requirements manually", height=150, 
                               disabled=st.session_state.file_uploaded)
    
    # If manual input is used, disable file uploader
    if manual_input:
        st.session_state.manual_input_used = True
    else:
        st.session_state.manual_input_used = False
    
    # Extract text from file or use manual input
    def extract_text(file):
        if not file: return ""
        try:
            if file.type == "text/plain":
                return file.getvalue().decode()
            elif file.type == "text/csv":
                return pd.read_csv(file).to_string()
            elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return "\n".join([p.text for p in Document(file).paragraphs])
        except Exception as e:
            st.error(f"File error: {e}")
            return ""
    
    requirements = extract_text(criteria_file) or manual_input

    # Add input for number of test cases
    num_test_cases = st.number_input("Number of Test Cases", min_value=1, max_value=20, value=5)

# Main interface
if st.button("✨ Generate Test Cases"):
    if not requirements:
        st.warning("Please input requirements first!")
    else:
        try:
            with st.spinner("Generating test cases..."):
                generated = generate_test_cases(insurance_type, region, line_of_business, requirements, num_test_cases)
                st.session_state.test_cases = generated
                st.success(f"Generated {len(generated)} test cases!")
        except Exception as e:
            st.error(f"Generation failed: {str(e)}")

st.divider()

# Display results with debugging
if 'test_cases' in st.session_state:
    st.subheader("Generated Test Cases")
    
    # Debug view
    with st.expander("Debug Raw Response"):
        st.code(st.session_state.raw_response)
    
    # Main table display
    if st.session_state.test_cases:
        df = pd.DataFrame(st.session_state.test_cases)
        
        # Remove LOB and Region columns
        df_display = df.drop(columns=["LOB", "Region"])
        
        # Format steps with line breaks
        df_display['Execution Steps'] = df_display['Execution Steps'].str.replace(r'(Step\d+:)', r'\n\1', regex=True)
        
        # Fix: Use raw string for the HTML replacement
        html_table = df_display.to_html(index=False, escape=False)
        html_table = html_table.replace(r'\n', '<br>')
        
        st.markdown(f"""
        <div style='overflow-x: auto; margin: 20px 0;'>
            {html_table}
        </div>
        """, unsafe_allow_html=True)
        
        # Excel export (without LOB and Region)
        excel_buffer = BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
            df_display.to_excel(writer, index=False)
            worksheet = writer.sheets['Sheet1']
            
            # Format columns
            for col in worksheet.columns:
                max_length = max(len(str(cell.value)) for cell in col)
                worksheet.column_dimensions[col[0].column_letter].width = max_length + 2
            
            # Highlight the header row in yellow
            yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
            for cell in worksheet[1]:  # Header row (first row)
                cell.fill = yellow_fill
            
            # Format steps
            for row in worksheet.iter_rows(min_row=2):
                cell = row[5]  # Execution Steps column (adjusted index after removing columns)
                cell.alignment = Alignment(wrap_text=True, vertical='top')

        st.download_button(
            "📥 Download Excel",
            excel_buffer.getvalue(),
            "test_cases.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        
        # TXT export (without LOB and Region)
        txt_content = format_test_cases_for_txt(st.session_state.test_cases)
        st.download_button(
            "📥 Download as TXT",
            txt_content,
            "test_cases.txt",
            "text/plain"
        )
        
        # Word (DOCX) export (without LOB and Region)
        doc = format_test_cases_for_docx(st.session_state.test_cases)
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        st.download_button(
            "📥 Download as Word (DOCX)",
            doc_buffer.getvalue(),
            "test_cases.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("No valid test cases could be parsed from the response")

elif 'test_cases' not in st.session_state:
    st.info("Click 'Generate Test Cases' to begin")